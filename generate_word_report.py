from pathlib import Path
from datetime import datetime

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

from docx import Document
from docx.shared import Inches


BASE_DIR = Path(__file__).resolve().parent
DATA_PATH = BASE_DIR / "data.csv"
FIG_DIR = BASE_DIR / "report_figures"
REPORT_PATH = BASE_DIR / "Rossmann_EDA_Report.docx"

FIG_DIR.mkdir(exist_ok=True)

sns.set_theme(style="whitegrid", context="notebook")
plt.rcParams["figure.figsize"] = (10, 5)


def add_heading_and_caption(doc: Document, title: str, caption: str):
    doc.add_heading(title, level=3)
    doc.add_paragraph(caption)


def add_dataframe_table(doc: Document, df: pd.DataFrame, title: str):
    doc.add_paragraph(title)
    df_to_show = df.copy()
    if isinstance(df_to_show.index, pd.MultiIndex) or df_to_show.index.name is not None:
        df_to_show = df_to_show.reset_index()

    table = doc.add_table(rows=1, cols=len(df_to_show.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_to_show.columns):
        hdr_cells[i].text = str(col)

    for _, row in df_to_show.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float):
                row_cells[i].text = f"{val:,.4f}" if abs(val) < 100 else f"{val:,.2f}"
            else:
                row_cells[i].text = str(val)


def main():
    if not DATA_PATH.exists():
        raise FileNotFoundError(f"Missing data file: {DATA_PATH}")

    df = pd.read_csv(DATA_PATH, low_memory=False)
    df.columns = [c.strip() for c in df.columns]
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")

    flag_map = {"yes": 1, "y": 1, "true": 1, "1": 1, "no": 0, "n": 0, "false": 0, "0": 0}
    binary_cols = ["IsOpen", "Open_Flag", "StoreClosed_Flag", "DailyPromo", "SchoolHoliday_Flag", "RecurringPromo", "CompetitorDistance_Imputed"]
    for c in binary_cols:
        if c in df.columns:
            norm = df[c].astype(str).str.strip().str.lower().map(flag_map)
            df[c] = pd.to_numeric(norm.fillna(df[c]), errors="coerce")

    df["Year"] = df["Date"].dt.year
    df["Month"] = df["Date"].dt.month
    df["MonthName"] = df["Date"].dt.month_name()
    df["Quarter"] = df["Date"].dt.quarter
    df["Week"] = df["Date"].dt.isocalendar().week.astype("Int64")
    df["SalesPerCustomer"] = np.where(df["CustomerCount"].fillna(0) > 0, df["Sales"] / df["CustomerCount"], np.nan)

    open_flag_col = "Open_Flag" if "Open_Flag" in df.columns else "IsOpen"
    open_df = df[(df["Sales"] > 0) & (df[open_flag_col] == 1)].copy()

    # Figure 1: Missingness
    missing_pct = (df.isna().mean() * 100).sort_values(ascending=False)
    plt.figure(figsize=(12, 5))
    missing_pct[missing_pct > 0].plot(kind="bar", color="#2E86AB")
    plt.title("Figure 1. Missingness by Column (%)")
    plt.ylabel("Missing %")
    plt.tight_layout()
    fig1 = FIG_DIR / "figure1_missingness.png"
    plt.savefig(fig1, dpi=180)
    plt.close()

    # Figure 2: Sales distributions
    fig2 = FIG_DIR / "figure2_sales_distribution.png"
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    sns.histplot(df["Sales"], bins=60, kde=True, ax=axes[0], color="#3D5A80")
    axes[0].set_title("All Sales")
    sns.histplot(np.log1p(open_df["Sales"]), bins=60, kde=True, ax=axes[1], color="#EE6C4D")
    axes[1].set_title("Log Sales (Open Days)")
    fig.suptitle("Figure 2. Sales Distribution")
    plt.tight_layout()
    plt.savefig(fig2, dpi=180)
    plt.close()

    # Figure 3: Sales vs Customers
    fig3 = FIG_DIR / "figure3_sales_vs_customers.png"
    sample_scatter = open_df.sample(min(len(open_df), 30000), random_state=42)
    plt.figure(figsize=(10, 5))
    sns.scatterplot(data=sample_scatter, x="CustomerCount", y="Sales", alpha=0.25, s=10)
    plt.title("Figure 3. Sales vs CustomerCount")
    plt.tight_layout()
    plt.savefig(fig3, dpi=180)
    plt.close()

    # Figure 4: Daily and Monthly trends
    fig4 = FIG_DIR / "figure4_time_trends.png"
    daily_sales = df.groupby("Date", as_index=False)["Sales"].sum()
    monthly_sales = df.set_index("Date").resample("ME")["Sales"].sum().reset_index()
    fig, axes = plt.subplots(2, 1, figsize=(14, 8))
    sns.lineplot(data=daily_sales, x="Date", y="Sales", ax=axes[0], color="#1B998B")
    axes[0].set_title("Daily Total Sales")
    sns.lineplot(data=monthly_sales, x="Date", y="Sales", ax=axes[1], color="#E84855")
    axes[1].set_title("Monthly Total Sales")
    fig.suptitle("Figure 4. Time Trends")
    plt.tight_layout()
    plt.savefig(fig4, dpi=180)
    plt.close()

    # Figure 5: Promo effect by store type
    fig5 = FIG_DIR / "figure5_promo_storetype.png"
    promo_by_store = open_df.groupby(["StoreType", "DailyPromo"], as_index=False)["Sales"].mean()
    plt.figure(figsize=(10, 5))
    sns.barplot(data=promo_by_store, x="StoreType", y="Sales", hue="DailyPromo")
    plt.title("Figure 5. Average Sales by StoreType and Promo")
    plt.tight_layout()
    plt.savefig(fig5, dpi=180)
    plt.close()

    # Figure 6: Correlation heatmap
    fig6 = FIG_DIR / "figure6_correlation_heatmap.png"
    corr_cols = [
        "Sales", "CustomerCount", "SalesPerCustomer", "DailyPromo", "SchoolHoliday_Flag",
        "RecurringPromo", "CompetitorDistance_m", "Open_Flag", "DayOfWeek", "Month", "Quarter"
    ]
    corr_cols = [c for c in corr_cols if c in df.columns]
    corr = df[corr_cols].apply(pd.to_numeric, errors="coerce").corr()
    plt.figure(figsize=(11, 8))
    sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm", center=0)
    plt.title("Figure 6. Correlation Matrix")
    plt.tight_layout()
    plt.savefig(fig6, dpi=180)
    plt.close()

    # Figure 7: Anomaly sensitivity
    fig7 = FIG_DIR / "figure7_anomaly_sensitivity.png"
    d = open_df.groupby("Date", as_index=False)["Sales"].sum().sort_values("Date")
    d["m28"] = d["Sales"].rolling(28, min_periods=14).mean()
    d["s28"] = d["Sales"].rolling(28, min_periods=14).std().replace(0, np.nan)
    d["z"] = (d["Sales"] - d["m28"]) / d["s28"]
    plt.figure(figsize=(13, 5))
    plt.plot(d["Date"], d["Sales"], label="Daily Sales", lw=1)
    for thr, color in [(2.0, "orange"), (2.5, "red")]:
        idx = d["z"].abs() >= thr
        plt.scatter(d.loc[idx, "Date"], d.loc[idx, "Sales"], s=12, color=color, label=f"|z|>={thr}")
    plt.title("Figure 7. Daily Sales with Anomaly Thresholds")
    plt.legend()
    plt.tight_layout()
    plt.savefig(fig7, dpi=180)
    plt.close()

    # Key tables
    numeric_desc = df.describe(include=[np.number]).T
    numeric_desc["median"] = df.select_dtypes(include=[np.number]).median()
    numeric_desc = numeric_desc[["count", "mean", "median", "std", "min", "25%", "50%", "75%", "max"]]

    columns_desc = {
        "StoreID": "Unique store identifier.",
        "Date": "Calendar date of observation.",
        "DayOfWeek": "Numeric day of week.",
        "DayName": "Day name label.",
        "Sales": "Total daily sales for the store.",
        "CustomerCount": "Total daily customers for the store.",
        "DailyPromo": "Indicator for whether daily promo is active.",
        "StateHoliday": "State holiday type/category.",
        "SchoolHoliday_Flag": "Indicator for school holiday.",
        "StoreType": "Store format/classification.",
        "Assortment": "Product assortment class.",
        "CompetitorDistance_m": "Distance to nearest competitor (meters).",
        "RecurringPromo": "Indicator for recurring promo participation.",
        "RecurringPromoStartWeek": "Promo2 start week.",
        "RecurringPromoStartYear": "Promo2 start year.",
        "RecurringPromoMonths": "Months with recurring promotion."
    }

    col_desc_rows = []
    for c in df.columns:
        col_desc_rows.append({
            "column": c,
            "description": columns_desc.get(c, "Derived or source column used in analysis."),
            "dtype": str(df[c].dtype)
        })
    col_desc_df = pd.DataFrame(col_desc_rows)

    range_rows = []
    for c in df.columns:
        s = df[c]
        dtype = str(s.dtype)
        if pd.api.types.is_numeric_dtype(s):
            rng = f"{s.min(skipna=True):,.2f} to {s.max(skipna=True):,.2f}"
        elif pd.api.types.is_datetime64_any_dtype(s):
            rng = f"{s.min(skipna=True)} to {s.max(skipna=True)}"
        else:
            nunique = s.nunique(dropna=True)
            rng = f"{nunique} unique values"
        notes = "Has missing values" if s.isna().any() else "No missing values"
        range_rows.append({"column": c, "dtype": dtype, "range_or_cardinality": rng, "notes": notes})
    ranges_df = pd.DataFrame(range_rows)

    promo_lift = open_df.groupby("DailyPromo")["Sales"].mean()
    promo_lift_pct = (promo_lift.loc[1] / promo_lift.loc[0] - 1) * 100 if (0 in promo_lift.index and 1 in promo_lift.index and promo_lift.loc[0] != 0) else np.nan
    corr_sales_customers = open_df[["Sales", "CustomerCount"]].corr().iloc[0, 1]

    monthly_rank = open_df.groupby(open_df["Date"].dt.month)["Sales"].mean().sort_values(ascending=False)
    strongest_month = int(monthly_rank.index[0])
    weakest_month = int(monthly_rank.index[-1])

    # Build Word document
    doc = Document()
    doc.add_heading("Rossmann Store Sales: Exploratory Data Analysis Report", 0)
    doc.add_paragraph(f"Prepared on {datetime.now():%Y-%m-%d}. This report presents dataset characteristics, preprocessing steps, exploratory findings, anomaly diagnostics, and proposed modeling methods.")

    doc.add_heading("1. Data Description", level=1)
    doc.add_paragraph("The analysis uses the merged Rossmann store-level daily sales data. The dataset contains calendar features, operational flags, promotional indicators, store metadata, and competition context variables.")

    add_dataframe_table(doc, col_desc_df, "Table 1. Data Column Descriptions")
    doc.add_paragraph("As shown in Table 1, the dataset includes core demand variables (Sales, CustomerCount), time indicators, and potential explanatory factors such as promotions, holidays, and store type.")

    add_dataframe_table(doc, ranges_df, "Table 2. Data Types, Ranges, and Notes")
    doc.add_paragraph("As shown in Table 2, several operational fields contain structured missingness, which is important for feature engineering decisions.")

    add_dataframe_table(doc, numeric_desc.round(3), "Table 3. Descriptive Statistics of Numeric Variables")
    doc.add_paragraph("Table 3 summarizes mean, median, dispersion, and quantiles, supporting baseline understanding of skewness and outlier potential.")

    doc.add_heading("2. Key Findings from Exploratory Data Analysis", level=1)

    add_heading_and_caption(doc, "Figure 1. Missingness by Column", "Figure 1 shows concentrated missingness in a subset of fields, indicating non-random, structural missingness. This supports indicator-aware imputation for selected variables.")
    doc.add_picture(str(fig1), width=Inches(6.8))

    add_heading_and_caption(doc, "Figure 2. Sales Distribution", "Figure 2 shows right-skewed sales distributions. The log transformation reduces skewness and improves stability for downstream models.")
    doc.add_picture(str(fig2), width=Inches(6.8))

    add_heading_and_caption(doc, "Figure 3. Sales vs CustomerCount", "Figure 3 indicates a strong positive relationship between customer traffic and sales, suggesting traffic is a dominant short-term sales driver.")
    doc.add_picture(str(fig3), width=Inches(6.8))

    add_heading_and_caption(doc, "Figure 4. Daily and Monthly Time Trends", "Figure 4 highlights nonstationary demand patterns and seasonality. Aggregate monthly trends are smoother and show systematic cyclical behavior.")
    doc.add_picture(str(fig4), width=Inches(6.8))

    add_heading_and_caption(doc, "Figure 5. Promo Effect by StoreType", "Figure 5 shows promotions raise average sales across store types, but effect size differs by segment, motivating interaction terms.")
    doc.add_picture(str(fig5), width=Inches(6.8))

    add_heading_and_caption(doc, "Figure 6. Correlation Matrix", "Figure 6 summarizes linear relationships. Sales and CustomerCount exhibit strong positive correlation, while promotion and temporal features provide additional explanatory signal.")
    doc.add_picture(str(fig6), width=Inches(6.8))

    add_heading_and_caption(doc, "Figure 7. Anomaly Sensitivity", "Figure 7 demonstrates anomaly detection sensitivity to threshold choice. Lower thresholds capture more events and are appropriate for monitoring workflows.")
    doc.add_picture(str(fig7), width=Inches(6.8))

    doc.add_paragraph(
        "EDA Summary: As shown across Figures 2-6, sales are driven primarily by customer traffic, with significant contribution from promotional activity and store segmentation factors. "
        f"Estimated promo uplift is {promo_lift_pct:,.2f}% on open days, and Sales-Customer correlation is {corr_sales_customers:,.3f}. "
        f"Seasonality is meaningful: month {strongest_month} has the highest average sales while month {weakest_month} has the lowest."
    )

    doc.add_heading("3. Proposed Methods or Models for Next Analysis", level=1)
    doc.add_paragraph("Based on the exploratory findings, the next stage should combine interpretable baselines and nonlinear methods:")
    doc.add_paragraph("1) Regularized Linear Regression (on log-sales): provides transparent coefficient-level interpretation and strong baseline stability.")
    doc.add_paragraph("2) Gradient-Boosted Trees / Random Forests: captures nonlinear effects and interactions such as promo by store type, and customer traffic saturation effects.")
    doc.add_paragraph("3) Time-aware validation and rolling-window backtesting: ensures realistic forecast evaluation and avoids temporal leakage.")
    doc.add_paragraph("4) Feature strategy: include lagged sales, calendar indicators, promo flags, store metadata, and missingness indicators for structurally absent fields.")

    doc.add_paragraph("As shown in Figure 5 and Figure 6, segment-specific and interaction-aware models are justified. As shown in Figure 7, anomaly-aware preprocessing should be incorporated into model training pipelines.")

    doc.add_heading("References", level=2)
    doc.add_paragraph("Hyndman, R. J., & Athanasopoulos, G. (2021). Forecasting: Principles and Practice (3rd ed.). OTexts.")
    doc.add_paragraph("James, G., Witten, D., Hastie, T., & Tibshirani, R. (2021). An Introduction to Statistical Learning (2nd ed.). Springer.")
    doc.add_paragraph("Lundberg, S. M., & Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions. NeurIPS.")
    doc.add_paragraph("Biecek, P., & Burzykowski, T. (2021). Explanatory Model Analysis. Chapman and Hall/CRC.")

    doc.save(REPORT_PATH)
    print(f"Report created: {REPORT_PATH}")


if __name__ == "__main__":
    main()
